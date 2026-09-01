from datetime import datetime, timezone
from dataclasses import dataclass
from fastapi import HTTPException
from sqlalchemy.orm import Session
from typing import Any, Iterator
from app.chats.export_security import (
    is_chat_excluded_from_default_export,
    sanitize_chat_share_for_export,
)
from app.chats.compliance import (
    ComplianceWatermarkResolver,
    append_compliance_watermark,
    apply_compliance_watermark_to_chat_export,
    get_compliance_watermark,
)
from app.chats.models import ChatReadState, Chats, ChatMessages
from app.chats.schemas import CHAT_IMPORT_MAX_DEEP_RESEARCH_BYTES_PER_CHAT
from app.files.utils import materialize_file_record, resolve_accessible_file_record
from app.users.models import User
from app.utils.svg import rasterize_svg_to_png_bytes
from app.utils.email import build_email_reference_token
from copy import deepcopy
from functools import lru_cache
from pathlib import Path
import base64
import mimetypes
import json
import logging


current_chat_export_version = 1.0
ALL_CHATS_EXPORT_QUERY_BATCH_SIZE = 200

ATTACHMENT_FIELDS = ("images", "videos", "audios", "documents")
DISPLAY_CONTENT_BLOCK_TYPES = {
    "content",
    "text",
    "message",
    "assistant",
    "user",
    "markdown",
}
DISPLAY_SKIP_BLOCK_TYPES = {"reasoning", "thinking", "internal", "meta"}
SUPPORTED_PDF_EXPORT_LANGUAGES = {
    "en",
    "de",
    "es",
    "zh",
    "fr",
    "hi",
    "ar",
    "ja",
    "it",
    "pt",
    "ru",
}
PDF_EXPORT_I18N_KEYS = {
    "pdf_export_title": "Chat export",
    "pdf_export_generated": "Generated {date}",
    "pdf_export_message_count_one": "{count} message",
    "pdf_export_message_count_other": "{count} messages",
    "pdf_export_started": "Started {date}",
    "pdf_export_updated": "Updated {date}",
    "pdf_export_user": "User",
    "pdf_export_assistant": "Assistant",
    "pdf_export_tool": "Tool",
    "pdf_export_tool_named": "Tool: {name}",
    "pdf_export_system": "System",
    "pdf_export_message": "Message",
    "pdf_export_tool_activity": "Tool activity was recorded for this turn.",
    "pdf_export_page": "Page {page}",
    "pdf_export_images": "Images",
    "pdf_export_videos": "Videos",
    "pdf_export_audios": "Audio",
    "pdf_export_documents": "Documents",
}

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class PdfExportAttachment:
    """Resolved file metadata used when rendering chat attachments into PDFs."""

    field: str
    file_id: str
    label: str
    file_type: str
    file_category: str
    file_size: int
    path: Path | None = None


def _json_dumps(value: Any) -> str:
    """Serialize compact JSON for streamed chat exports."""
    return json.dumps(value, ensure_ascii=True, default=str, separators=(",", ":"))


def _iter_export_query_rows(query, batch_size: int = ALL_CHATS_EXPORT_QUERY_BATCH_SIZE):
    """Iterate query results in batches instead of materializing full lists."""
    if hasattr(query, "execution_options"):
        query = query.execution_options(stream_results=True)
    if hasattr(query, "yield_per"):
        query = query.yield_per(batch_size)
    try:
        yield from query
    except TypeError:
        yield from query.all()


PDF_EXPORT_TRANSLATIONS = {
    "en": PDF_EXPORT_I18N_KEYS,
    "de": {
        "pdf_export_title": "Chat-Export",
        "pdf_export_generated": "Erstellt am {date}",
        "pdf_export_message_count_one": "{count} Nachricht",
        "pdf_export_message_count_other": "{count} Nachrichten",
        "pdf_export_started": "Gestartet am {date}",
        "pdf_export_updated": "Aktualisiert am {date}",
        "pdf_export_user": "Benutzer",
        "pdf_export_assistant": "Assistent",
        "pdf_export_tool": "Tool",
        "pdf_export_tool_named": "Tool: {name}",
        "pdf_export_system": "System",
        "pdf_export_message": "Nachricht",
        "pdf_export_tool_activity": "Für diese Runde wurde eine Tool-Aktivität aufgezeichnet.",
        "pdf_export_page": "Seite {page}",
        "pdf_export_images": "Bilder",
        "pdf_export_videos": "Videos",
        "pdf_export_audios": "Audio",
        "pdf_export_documents": "Dokumente",
    },
    "es": {
        "pdf_export_title": "Exportación del chat",
        "pdf_export_generated": "Generado el {date}",
        "pdf_export_message_count_one": "{count} mensaje",
        "pdf_export_message_count_other": "{count} mensajes",
        "pdf_export_started": "Iniciado el {date}",
        "pdf_export_updated": "Actualizado el {date}",
        "pdf_export_user": "Usuario",
        "pdf_export_assistant": "Asistente",
        "pdf_export_tool": "Herramienta",
        "pdf_export_tool_named": "Herramienta: {name}",
        "pdf_export_system": "Sistema",
        "pdf_export_message": "Mensaje",
        "pdf_export_tool_activity": "Se registró actividad de herramienta en este turno.",
        "pdf_export_page": "Página {page}",
        "pdf_export_images": "Imágenes",
        "pdf_export_videos": "Videos",
        "pdf_export_audios": "Audios",
        "pdf_export_documents": "Documentos",
    },
    "zh": {
        "pdf_export_title": "聊天导出",
        "pdf_export_generated": "生成于 {date}",
        "pdf_export_message_count_one": "{count} 条消息",
        "pdf_export_message_count_other": "{count} 条消息",
        "pdf_export_started": "开始于 {date}",
        "pdf_export_updated": "更新于 {date}",
        "pdf_export_user": "用户",
        "pdf_export_assistant": "助手",
        "pdf_export_tool": "工具",
        "pdf_export_tool_named": "工具：{name}",
        "pdf_export_system": "系统",
        "pdf_export_message": "消息",
        "pdf_export_tool_activity": "本轮记录了工具活动。",
        "pdf_export_page": "第 {page} 页",
        "pdf_export_images": "图片",
        "pdf_export_videos": "视频",
        "pdf_export_audios": "音频",
        "pdf_export_documents": "文档",
    },
    "fr": {
        "pdf_export_title": "Export du chat",
        "pdf_export_generated": "Généré le {date}",
        "pdf_export_message_count_one": "{count} message",
        "pdf_export_message_count_other": "{count} messages",
        "pdf_export_started": "Démarré le {date}",
        "pdf_export_updated": "Mis à jour le {date}",
        "pdf_export_user": "Utilisateur",
        "pdf_export_assistant": "Assistant",
        "pdf_export_tool": "Outil",
        "pdf_export_tool_named": "Outil : {name}",
        "pdf_export_system": "Système",
        "pdf_export_message": "Message",
        "pdf_export_tool_activity": "Une activité d’outil a été enregistrée pour ce tour.",
        "pdf_export_page": "Page {page}",
        "pdf_export_images": "Images",
        "pdf_export_videos": "Vidéos",
        "pdf_export_audios": "Audio",
        "pdf_export_documents": "Documents",
    },
    "hi": {
        "pdf_export_title": "चैट निर्यात",
        "pdf_export_generated": "{date} को जनरेट किया गया",
        "pdf_export_message_count_one": "{count} संदेश",
        "pdf_export_message_count_other": "{count} संदेश",
        "pdf_export_started": "{date} को शुरू हुआ",
        "pdf_export_updated": "{date} को अपडेट हुआ",
        "pdf_export_user": "उपयोगकर्ता",
        "pdf_export_assistant": "सहायक",
        "pdf_export_tool": "टूल",
        "pdf_export_tool_named": "टूल: {name}",
        "pdf_export_system": "सिस्टम",
        "pdf_export_message": "संदेश",
        "pdf_export_tool_activity": "इस चरण के लिए टूल गतिविधि दर्ज की गई।",
        "pdf_export_page": "पृष्ठ {page}",
        "pdf_export_images": "चित्र",
        "pdf_export_videos": "वीडियो",
        "pdf_export_audios": "ऑडियो",
        "pdf_export_documents": "दस्तावेज़",
    },
    "ar": {
        "pdf_export_title": "تصدير المحادثة",
        "pdf_export_generated": "تم الإنشاء في {date}",
        "pdf_export_message_count_one": "{count} رسالة",
        "pdf_export_message_count_other": "{count} رسائل",
        "pdf_export_started": "بدأ في {date}",
        "pdf_export_updated": "تم التحديث في {date}",
        "pdf_export_user": "المستخدم",
        "pdf_export_assistant": "المساعد",
        "pdf_export_tool": "أداة",
        "pdf_export_tool_named": "أداة: {name}",
        "pdf_export_system": "النظام",
        "pdf_export_message": "رسالة",
        "pdf_export_tool_activity": "تم تسجيل نشاط أداة لهذه الجولة.",
        "pdf_export_page": "الصفحة {page}",
        "pdf_export_images": "الصور",
        "pdf_export_videos": "الفيديوهات",
        "pdf_export_audios": "الصوت",
        "pdf_export_documents": "المستندات",
    },
    "ja": {
        "pdf_export_title": "チャットのエクスポート",
        "pdf_export_generated": "{date} に生成",
        "pdf_export_message_count_one": "{count} 件のメッセージ",
        "pdf_export_message_count_other": "{count} 件のメッセージ",
        "pdf_export_started": "{date} に開始",
        "pdf_export_updated": "{date} に更新",
        "pdf_export_user": "ユーザー",
        "pdf_export_assistant": "アシスタント",
        "pdf_export_tool": "ツール",
        "pdf_export_tool_named": "ツール: {name}",
        "pdf_export_system": "システム",
        "pdf_export_message": "メッセージ",
        "pdf_export_tool_activity": "このターンでツールのアクティビティが記録されました。",
        "pdf_export_page": "ページ {page}",
        "pdf_export_images": "画像",
        "pdf_export_videos": "動画",
        "pdf_export_audios": "音声",
        "pdf_export_documents": "ドキュメント",
    },
    "it": {
        "pdf_export_title": "Esportazione chat",
        "pdf_export_generated": "Generato il {date}",
        "pdf_export_message_count_one": "{count} messaggio",
        "pdf_export_message_count_other": "{count} messaggi",
        "pdf_export_started": "Avviato il {date}",
        "pdf_export_updated": "Aggiornato il {date}",
        "pdf_export_user": "Utente",
        "pdf_export_assistant": "Assistente",
        "pdf_export_tool": "Strumento",
        "pdf_export_tool_named": "Strumento: {name}",
        "pdf_export_system": "Sistema",
        "pdf_export_message": "Messaggio",
        "pdf_export_tool_activity": "È stata registrata attività dello strumento per questo turno.",
        "pdf_export_page": "Pagina {page}",
        "pdf_export_images": "Immagini",
        "pdf_export_videos": "Video",
        "pdf_export_audios": "Audio",
        "pdf_export_documents": "Documenti",
    },
    "pt": {
        "pdf_export_title": "Exportação do chat",
        "pdf_export_generated": "Gerado em {date}",
        "pdf_export_message_count_one": "{count} mensagem",
        "pdf_export_message_count_other": "{count} mensagens",
        "pdf_export_started": "Iniciado em {date}",
        "pdf_export_updated": "Atualizado em {date}",
        "pdf_export_user": "Usuário",
        "pdf_export_assistant": "Assistente",
        "pdf_export_tool": "Ferramenta",
        "pdf_export_tool_named": "Ferramenta: {name}",
        "pdf_export_system": "Sistema",
        "pdf_export_message": "Mensagem",
        "pdf_export_tool_activity": "A atividade da ferramenta foi registrada neste turno.",
        "pdf_export_page": "Página {page}",
        "pdf_export_images": "Imagens",
        "pdf_export_videos": "Vídeos",
        "pdf_export_audios": "Áudios",
        "pdf_export_documents": "Documentos",
    },
    "ru": {
        "pdf_export_title": "Экспорт чата",
        "pdf_export_generated": "Создано {date}",
        "pdf_export_message_count_one": "{count} сообщение",
        "pdf_export_message_count_other": "{count} сообщений",
        "pdf_export_started": "Начато {date}",
        "pdf_export_updated": "Обновлено {date}",
        "pdf_export_user": "Пользователь",
        "pdf_export_assistant": "Ассистент",
        "pdf_export_tool": "Инструмент",
        "pdf_export_tool_named": "Инструмент: {name}",
        "pdf_export_system": "Система",
        "pdf_export_message": "Сообщение",
        "pdf_export_tool_activity": "Для этого шага была записана активность инструмента.",
        "pdf_export_page": "Страница {page}",
        "pdf_export_images": "Изображения",
        "pdf_export_videos": "Видео",
        "pdf_export_audios": "Аудио",
        "pdf_export_documents": "Документы",
    },
}


def _decode_json_content(raw_content):
    """Attempt to decode raw content (str, list, dict, or None) into a parsed JSON structure.

    Returns None for empty or invalid inputs.
    """
    if raw_content is None:
        return None
    if isinstance(raw_content, (list, dict)):
        return raw_content
    if isinstance(raw_content, str):
        stripped = raw_content.strip()
        if not stripped:
            return None
        try:
            return json.loads(stripped)
        except Exception:
            return None
    return None


def _parse_attachment_values(raw_value) -> list[str]:
    """Parse raw attachment values into a deduplicated list of file ID strings."""
    if raw_value is None:
        return []

    value = raw_value
    if isinstance(value, str):
        stripped = value.strip()
        if not stripped:
            return []
        try:
            value = json.loads(stripped)
        except Exception:
            return [stripped]

    if isinstance(value, list):
        out: list[str] = []
        seen: set[str] = set()
        for item in value:
            file_id = (
                item.get("id") or item.get("file_id")
                if isinstance(item, dict)
                else item
            )
            normalized = str(file_id or "").strip()
            if not normalized or normalized in seen:
                continue
            seen.add(normalized)
            out.append(normalized)
        return out

    normalized = str(value).strip()
    return [normalized] if normalized else []


def _collect_message_attachments(message_row) -> dict[str, list[str] | None]:
    """Collect all attachment IDs (images, videos, audios, documents) from a message row.

    Includes attachments embedded in JSON content blocks.
    """
    collected = {field: [] for field in ATTACHMENT_FIELDS}
    seen = {field: set() for field in ATTACHMENT_FIELDS}

    def _append(field: str, values: list[str]) -> None:
        for value in values:
            if value in seen[field]:
                continue
            seen[field].add(value)
            collected[field].append(value)

    for field in ATTACHMENT_FIELDS:
        _append(field, _parse_attachment_values(getattr(message_row, field, None)))

    decoded_content = _decode_json_content(getattr(message_row, "content", None))
    blocks = (
        decoded_content
        if isinstance(decoded_content, list)
        else [decoded_content]
        if isinstance(decoded_content, dict)
        else []
    )
    for block in blocks:
        if not isinstance(block, dict):
            continue
        for field in ATTACHMENT_FIELDS:
            _append(field, _parse_attachment_values(block.get(field)))

    return {field: (collected[field] or None) for field in ATTACHMENT_FIELDS}


def _coerce_display_text(value) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, (int, float, bool)):
        return str(value)
    return json.dumps(value, ensure_ascii=False, indent=2)


def _display_content_from_block(block: dict) -> str:
    block_type = str(block.get("type") or "").strip().lower()
    if block_type in DISPLAY_SKIP_BLOCK_TYPES:
        return ""
    if block_type and block_type not in DISPLAY_CONTENT_BLOCK_TYPES:
        return ""
    for key in ("content", "text", "value", "markdown"):
        text = _coerce_display_text(block.get(key))
        if text:
            return text
    return ""


def _message_display_content(message: dict) -> str:
    """Return user-visible message content, hiding stored block/meta payloads."""
    raw_content = message.get("content")
    decoded_content = _decode_json_content(raw_content)
    blocks: list = []
    if isinstance(decoded_content, list):
        blocks = decoded_content
    elif isinstance(decoded_content, dict):
        blocks = [decoded_content]

    if blocks:
        parts: list[str] = []
        for block in blocks:
            if isinstance(block, dict):
                text = _display_content_from_block(block)
            else:
                text = _coerce_display_text(block)
            if text:
                parts.append(text)
        return "\n\n".join(parts).strip()

    return _coerce_display_text(raw_content)


def _attachment_label(item) -> str:
    if isinstance(item, dict):
        meta = item.get("meta") if isinstance(item.get("meta"), dict) else {}
        label = (
            item.get("name")
            or item.get("filename")
            or item.get("original_filename")
            or meta.get("original_filename")
            or item.get("id")
            or item.get("file_id")
        )
        return str(label or "").strip()
    return str(item or "").strip()


def _parse_attachment_labels(raw_value) -> list[str]:
    if raw_value is None:
        return []
    value = raw_value
    if isinstance(value, str):
        stripped = value.strip()
        if not stripped:
            return []
        try:
            value = json.loads(stripped)
        except Exception:
            return [stripped]
    if isinstance(value, list):
        labels = [_attachment_label(item) for item in value]
    else:
        labels = [_attachment_label(value)]
    out: list[str] = []
    seen: set[str] = set()
    for label in labels:
        if not label or label in seen:
            continue
        seen.add(label)
        out.append(label)
    return out


def _message_attachment_labels(message: dict) -> dict[str, list[str]]:
    out: dict[str, list[str]] = {}

    def _extend(field: str, labels: list[str]) -> None:
        if not labels:
            return
        existing = out.setdefault(field, [])
        seen = set(existing)
        for label in labels:
            if label in seen:
                continue
            seen.add(label)
            existing.append(label)

    for field in ATTACHMENT_FIELDS:
        _extend(field, _parse_attachment_labels(message.get(field)))

    decoded_content = _decode_json_content(message.get("content"))
    blocks = (
        decoded_content
        if isinstance(decoded_content, list)
        else [decoded_content]
        if isinstance(decoded_content, dict)
        else []
    )
    for block in blocks:
        if not isinstance(block, dict):
            continue
        for field in ATTACHMENT_FIELDS:
            _extend(field, _parse_attachment_labels(block.get(field)))

    return out


def _message_attachment_ids(message: dict) -> dict[str, list[str]]:
    """Return attachment file IDs from message fields and embedded content blocks."""
    out: dict[str, list[str]] = {}

    def _extend(field: str, file_ids: list[str]) -> None:
        if not file_ids:
            return
        existing = out.setdefault(field, [])
        seen = set(existing)
        for file_id in file_ids:
            if file_id in seen:
                continue
            seen.add(file_id)
            existing.append(file_id)

    for field in ATTACHMENT_FIELDS:
        _extend(field, _parse_attachment_values(message.get(field)))

    decoded_content = _decode_json_content(message.get("content"))
    blocks = (
        decoded_content
        if isinstance(decoded_content, list)
        else [decoded_content]
        if isinstance(decoded_content, dict)
        else []
    )
    for block in blocks:
        if not isinstance(block, dict):
            continue
        for field in ATTACHMENT_FIELDS:
            _extend(field, _parse_attachment_values(block.get(field)))

    return out


def _format_pdf_file_size(byte_count: int | None) -> str:
    """Format a file size using the same compact units as the chat file tiles."""
    size = max(0, int(byte_count or 0))
    if size == 0:
        return "0 B"
    units = ("B", "KB", "MB", "GB")
    value = float(size)
    unit_index = 0
    while value >= 1024 and unit_index < len(units) - 1:
        value /= 1024
        unit_index += 1
    return f"{value:.2f} {units[unit_index]}"


def _pdf_file_extension_label(filename: str, file_type: str) -> str:
    """Return a short extension label for PDF file-card badges."""
    suffix = Path(str(filename or "")).suffix.lstrip(".")
    if suffix:
        return suffix[:8].upper()
    mime_suffix = str(file_type or "").split("/", 1)[-1].strip()
    return (mime_suffix or "FILE")[:8].upper()


def _pdf_attachment_filename(
    file_record, fallback_label: str, fallback_file_id: str
) -> str:
    """Resolve the display filename used for a PDF attachment card or image caption."""
    meta = getattr(file_record, "meta", None) if file_record is not None else None
    meta = meta if isinstance(meta, dict) else {}
    label = (
        meta.get("original_filename")
        or meta.get("original_name")
        or getattr(file_record, "file_name", None)
        or fallback_label
        or fallback_file_id
    )
    return (
        Path(str(label or fallback_file_id or "file").replace("\x00", "")).name
        or "file"
    )


def _resolve_pdf_export_attachments(
    message: dict, db, user_id: str
) -> dict[str, list[PdfExportAttachment]]:
    """Resolve exported attachment IDs into metadata that the PDF renderer can display."""
    attachment_ids = _message_attachment_ids(message)
    attachment_labels = _message_attachment_labels(message)
    resolved: dict[str, list[PdfExportAttachment]] = {}

    for field, file_ids in attachment_ids.items():
        field_attachments: list[PdfExportAttachment] = []
        labels = attachment_labels.get(field, [])
        for index, file_id in enumerate(file_ids):
            fallback_label = labels[index] if index < len(labels) else file_id
            file_record = None
            owner_user_id = None
            try:
                file_record, owner_user_id = resolve_accessible_file_record(
                    db, str(user_id), str(file_id)
                )
            except Exception:
                logger.debug(
                    "Unable to resolve chat PDF export attachment metadata",
                    exc_info=True,
                )

            if not file_record or not owner_user_id:
                field_attachments.append(
                    PdfExportAttachment(
                        field=field,
                        file_id=str(file_id),
                        label=str(fallback_label or file_id),
                        file_type="",
                        file_category=field[:-1] if field.endswith("s") else field,
                        file_size=0,
                        path=None,
                    )
                )
                continue

            file_type = str(getattr(file_record, "file_type", "") or "").strip().lower()
            file_category = (
                str(getattr(file_record, "file_category", "") or "").strip().lower()
            )
            file_size = int(getattr(file_record, "file_size", 0) or 0)
            label = _pdf_attachment_filename(
                file_record, str(fallback_label or ""), str(file_id)
            )
            path = None
            if file_type.startswith("image/"):
                try:
                    path = materialize_file_record(file_record, str(owner_user_id))
                except Exception:
                    logger.debug(
                        "Unable to materialize chat PDF export image attachment",
                        exc_info=True,
                    )

            field_attachments.append(
                PdfExportAttachment(
                    field=field,
                    file_id=str(file_id),
                    label=label,
                    file_type=file_type,
                    file_category=file_category,
                    file_size=file_size,
                    path=path,
                )
            )

        if field_attachments:
            resolved[field] = field_attachments

    return resolved


def _format_export_datetime(value: str | None) -> str:
    if not value:
        return ""
    normalized = str(value).strip()
    if not normalized:
        return ""
    try:
        parsed = datetime.fromisoformat(normalized.replace("Z", "+00:00"))
    except ValueError:
        return normalized
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=timezone.utc)
    parsed = parsed.astimezone(timezone.utc)
    return parsed.strftime("%Y-%m-%d %H:%M UTC")


def _resolve_application_name(db) -> str:
    try:
        from app.settings.utils import get_value_by_page_and_key

        value = get_value_by_page_and_key("general", "application_name", db)
        if isinstance(value, str) and value.strip():
            return value.strip()
    except Exception:
        pass
    return "Omlorix"


def _normalize_pdf_export_language(value) -> str:
    if not isinstance(value, str):
        return "en"
    normalized = value.strip().lower().replace("_", "-")
    if not normalized:
        return "en"
    primary = normalized.split("-", 1)[0]
    return primary if primary in SUPPORTED_PDF_EXPORT_LANGUAGES else "en"


def _language_from_user_settings(settings) -> str:
    if not isinstance(settings, dict):
        return ""
    general = settings.get("general")
    if not isinstance(general, dict):
        return ""
    language = general.get("language")
    return language.strip() if isinstance(language, str) else ""


def _resolve_pdf_export_language(user_id: str, db) -> str:
    language = ""
    try:
        user = db.query(User).filter(User.id == user_id).first()
        language = _language_from_user_settings(getattr(user, "settings", None))
    except Exception:
        language = ""
    if language:
        return _normalize_pdf_export_language(language)

    try:
        from app.users.init import get_user_setting_value

        language = get_user_setting_value(user_id, "general", "language", db)
        return _normalize_pdf_export_language(language)
    except Exception:
        return "en"


@lru_cache(maxsize=32)
def _load_pdf_export_translations(language: str) -> dict[str, str]:
    language = _normalize_pdf_export_language(language)
    return dict(PDF_EXPORT_TRANSLATIONS.get(language) or {})


def _pdf_export_t(language: str, key: str, **values) -> str:
    language = _normalize_pdf_export_language(language)
    fallback = PDF_EXPORT_I18N_KEYS.get(key, key)
    template = _load_pdf_export_translations(language).get(key)
    if not template and language != "en":
        template = _load_pdf_export_translations("en").get(key)
    template = template or fallback
    try:
        return template.format(**values)
    except Exception:
        return fallback.format(**values) if values else fallback


def _message_role_label(language: str, role: str, name: str | None = None) -> str:
    role = (role or "").lower()
    if role == "user":
        return _pdf_export_t(language, "pdf_export_user")
    if role == "assistant":
        return _pdf_export_t(language, "pdf_export_assistant")
    if role == "tool":
        return (
            _pdf_export_t(language, "pdf_export_tool_named", name=name)
            if name
            else _pdf_export_t(language, "pdf_export_tool")
        )
    if role == "system":
        return _pdf_export_t(language, "pdf_export_system")
    return role.capitalize() or _pdf_export_t(language, "pdf_export_message")


def _message_display_text_or_placeholder(message: dict, language: str) -> str:
    content = _message_display_content(message)
    if (message.get("role") or "").lower() == "tool" and not content:
        return _pdf_export_t(language, "pdf_export_tool_activity")
    return content


def _message_heading(message: dict, language: str) -> str:
    label = _message_role_label(
        language, str(message.get("role") or ""), message.get("name")
    )
    created_at = _format_export_datetime(message.get("created_at"))
    return f"{label} - {created_at}" if created_at else label


def _attachment_heading(language: str, field: str) -> str:
    return _pdf_export_t(language, f"pdf_export_{field}")


def _chat_export_details(chat: dict, message_count: int, language: str) -> list[str]:
    created = _format_export_datetime(chat.get("created_at"))
    updated = _format_export_datetime(chat.get("last_updated_at"))
    count_key = (
        "pdf_export_message_count_one"
        if message_count == 1
        else "pdf_export_message_count_other"
    )
    details = [_pdf_export_t(language, count_key, count=message_count)]
    if created:
        details.append(_pdf_export_t(language, "pdf_export_started", date=created))
    if updated:
        details.append(_pdf_export_t(language, "pdf_export_updated", date=updated))
    return details


def _chat_export_header(
    chat: dict, messages: list[dict], app_name: str, language: str, chat_title: str
) -> list[str]:
    exported_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    return [
        app_name,
        f"{_pdf_export_t(language, 'pdf_export_title')} - {_pdf_export_t(language, 'pdf_export_generated', date=exported_at)}",
        chat_title,
        " | ".join(_chat_export_details(chat, len(messages), language)),
    ]


def _build_display_transcript(
    chat: dict, messages: list[dict], app_name: str, language: str, chat_title: str
) -> dict:
    header = _chat_export_header(chat, messages, app_name, language, chat_title)
    return {
        "application_name": app_name,
        "title": chat_title,
        "export_label": header[1],
        "summary": header[3],
        "messages": [
            {
                "role": str(message.get("role") or ""),
                "role_label": _message_role_label(
                    language, str(message.get("role") or ""), message.get("name")
                ),
                "heading": _message_heading(message, language),
                "content": _message_display_text_or_placeholder(message, language),
                "attachments": _message_attachment_labels(message),
                "created_at": message.get("created_at"),
            }
            for message in messages
        ],
    }


def _render_attachment_lines(message: dict, language: str) -> list[str]:
    lines: list[str] = []
    for field, labels in _message_attachment_labels(message).items():
        lines.append(f"{_attachment_heading(language, field)}: {', '.join(labels)}")
    return lines


@lru_cache(maxsize=16)
def _register_pdf_export_font(language: str) -> str:
    language = _normalize_pdf_export_language(language)
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfbase.ttfonts import TTFont

        if language == "zh":
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            return "STSong-Light"
        if language == "ja":
            pdfmetrics.registerFont(UnicodeCIDFont("HeiseiKakuGo-W5"))
            return "HeiseiKakuGo-W5"

        language_candidates = {
            "ar": [
                "/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf",
                "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
            ],
            "hi": [
                "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf",
                "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
            ],
            "ru": [
                "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
            ],
        }
        candidates = language_candidates.get(language, []) + [
            "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        ]
        for candidate in candidates:
            path = Path(candidate)
            if not path.exists() or not path.is_file():
                continue
            font_name = f"OmlorixPdf-{language}-{path.stem}".replace(" ", "")
            pdfmetrics.registerFont(TTFont(font_name, str(path)))
            return font_name
    except Exception:
        pass
    return "Helvetica"


def _resolve_pdf_icon_path() -> Path | None:
    try:
        from app.settings import utils as settings_utils

        logo_dir = getattr(settings_utils, "_LOGO_DIR", None)
        candidates = [
            getattr(settings_utils, "_ICON_PNG_PATH", None),
            logo_dir / "favicon-512x512.png" if logo_dir else None,
            logo_dir / "apple-touch-icon.png" if logo_dir else None,
            logo_dir / "favicon-32x32.png" if logo_dir else None,
            logo_dir / "favicon-16x16.png" if logo_dir else None,
            logo_dir / "logo_light.png" if logo_dir else None,
            logo_dir / "logo_light.jpg" if logo_dir else None,
            logo_dir / "logo_light.jpeg" if logo_dir else None,
            logo_dir / "logo_light.webp" if logo_dir else None,
            getattr(settings_utils, "_FAVICON_SVG_PATH", None),
            logo_dir / "logo_light.svg" if logo_dir else None,
        ]
        for candidate in candidates:
            path = Path(candidate) if candidate else None
            if path and path.exists() and path.is_file():
                return path
    except Exception:
        pass
    return None


# -------------------
# Prepare Chat Download
# -------------------
def prepare_chat_download(
    user_id: str, chat_id: str, fmt: str, db, include_deleted_or_temp: bool = False
):
    """
    Prepare downloadable content for a chat.

    Args:
      user_id: Owner of the chat (enforced)
      chat_id: Chat to export
      fmt: "json" or "txt"
      db: session

    Returns:
      {"filename": str, "type": "json"|"txt", "content": dict|str}
    """
    data = export_chat_full(
        user_id, chat_id, db, include_deleted_or_temp=include_deleted_or_temp
    )
    compliance_watermark = get_compliance_watermark(user_id, db)

    # Helper to normalize attachments fields that may be stored as JSON strings
    def _parse_listish(val) -> list[str]:
        if not val:
            return []
        if isinstance(val, list):
            return [str(x) for x in val if str(x)]
        if isinstance(val, str):
            try:
                parsed = json.loads(val)
                if isinstance(parsed, list):
                    return [str(x) for x in parsed if str(x)]
            except Exception:
                pass
            s = val.strip()
            return [s] if s else []
        try:
            parsed = json.loads(val)
            if isinstance(parsed, list):
                return [str(x) for x in parsed if str(x)]
        except Exception:
            pass
        return []

    def _attachments_for_message(m: dict) -> dict[str, list[str]]:
        out = {}
        for field in ("images", "videos", "audios", "documents"):
            lst = _parse_listish(m.get(field))
            if lst:
                out[field] = lst
        return out

    def _sanitize(name: str) -> str:
        keep = [c for c in name if c.isalnum() or c in ("-", "_", " ")]
        s = "".join(keep).strip()
        return s or "chat"

    title = (
        data.get("chat", {}).get("title") or data.get("chat", {}).get("id") or "chat"
    )
    base = _sanitize(str(title))
    chat = data.get("chat", {})
    msgs = data.get("messages", []) or []
    app_name = _resolve_application_name(db)
    export_language = _resolve_pdf_export_language(user_id, db)
    chat_title = str(chat.get("title") or base or "Chat").strip() or "Chat"

    f = (fmt or "json").lower()
    if f == "json":
        cleaned = deepcopy(data)
        apply_compliance_watermark_to_chat_export(cleaned, compliance_watermark)
        cleaned["exported_at"] = datetime.now(timezone.utc).isoformat()
        cleaned["application_name"] = app_name
        cleaned["language"] = export_language
        cleaned["transcript"] = _build_display_transcript(
            chat,
            cleaned.get("messages", []) or [],
            app_name,
            export_language,
            chat_title,
        )
        msgs = cleaned.get("messages", []) or []
        for m in msgs:
            m["display_content"] = _message_display_content(m)
            attachment_labels = _message_attachment_labels(m)
            if attachment_labels:
                m["attachment_labels"] = attachment_labels
            # remove empty attachment keys; set non-empty to lists
            attach = _attachments_for_message(m)
            for field in ("images", "videos", "audios", "documents"):
                if field in m:
                    del m[field]
            for k, v in attach.items():
                m[k] = v
        return {"filename": f"{base}.json", "type": "json", "content": cleaned}
    if f == "txt":
        lines: list[str] = []
        header = _chat_export_header(chat, msgs, app_name, export_language, chat_title)
        lines.extend([header[0], header[1], "", header[2], header[3], ""])
        separator = "-" * 72
        for m in msgs:
            content = _message_display_text_or_placeholder(m, export_language)
            lines.append(separator)
            lines.append(_message_heading(m, export_language))
            lines.append("")
            if content:
                lines.append(content)
                lines.append("")
            attachment_lines = _render_attachment_lines(m, export_language)
            for attachment_line in attachment_lines:
                lines.append(attachment_line)
            if attachment_lines:
                lines.append("")
        text = append_compliance_watermark(
            "\n".join(lines).rstrip() + "\n", compliance_watermark
        )
        return {"filename": f"{base}.txt", "type": "txt", "content": text}

    if f == "md":
        import re as _re

        md_lines: list[str] = []

        def _demote_headings_by_two(md: str) -> str:
            # Add two # to any markdown heading line, capped at 6 total
            out_lines = []
            for ln in (md or "").splitlines():
                m = _re.match(r"^(#{1,6})\s+(.*)$", ln)
                if m:
                    cur = len(m.group(1))
                    new = min(6, cur + 2)
                    out_lines.append("#" * new + " " + m.group(2))
                else:
                    out_lines.append(ln)
            return "\n".join(out_lines)

        header = _chat_export_header(chat, msgs, app_name, export_language, chat_title)
        md_lines.append(f"# {header[2]}")
        md_lines.append("")
        md_lines.append(f"**{header[0]}**")
        md_lines.append("")
        md_lines.append(f"{header[1]}  ")
        md_lines.append(header[3])
        md_lines.append("")
        for m in msgs:
            content = _message_display_text_or_placeholder(m, export_language)
            md_lines.append(f"## {_message_heading(m, export_language)}")
            md_lines.append("")
            demoted = _demote_headings_by_two(content)
            if demoted:
                md_lines.append(demoted)
                md_lines.append("")
            attachment_lines = _render_attachment_lines(m, export_language)
            if attachment_lines:
                md_lines.append(
                    "**"
                    + " / ".join(
                        _attachment_heading(export_language, field)
                        for field in _message_attachment_labels(m)
                    )
                    + "**"
                )
                for attachment_line in attachment_lines:
                    md_lines.append(f"- {attachment_line}")
                md_lines.append("")

        md_text = append_compliance_watermark(
            "\n".join(md_lines).rstrip() + "\n", compliance_watermark
        )
        return {"filename": f"{base}.md", "type": "md", "content": md_text}

    if f == "docx":
        # Build a Word document export using python-docx
        try:
            from io import BytesIO
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from docx.oxml.ns import qn
        except Exception:
            raise HTTPException(
                status_code=501,
                detail="DOCX export requires the 'python-docx' package. Please install it (e.g., pip install python-docx).",
            )

        import re as _re

        doc = Document()
        # Basic style tweaks
        styles = doc.styles
        if "Normal" in styles:
            styles["Normal"].font.size = Pt(11)
            styles["Normal"].font.name = "Calibri"

        def _soft_wrap_long_tokens(
            text: str, max_token_len: int = 60, insert_every: int = 12
        ) -> str:
            ZWSP = "\u200b"
            out_lines = []
            for line in str(text).split("\n"):
                parts = []
                for tok in _re.split(r"(\s+)", line):
                    if not tok or tok.isspace():
                        parts.append(tok)
                        continue
                    if len(tok) > max_token_len:
                        chunks = [
                            tok[i : i + insert_every]
                            for i in range(0, len(tok), insert_every)
                        ]
                        parts.append(ZWSP.join(chunks))
                    else:
                        parts.append(tok)
                out_lines.append("".join(parts))
            return "\n".join(out_lines)

        def _add_title(text: str):
            p = doc.add_paragraph()
            run = p.add_run(" ".join(text.splitlines()).strip() or "Chat")
            run.font.size = Pt(20)
            run.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            doc.add_paragraph("")

        def _add_meta(text: str):
            p = doc.add_paragraph(text)
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.color.rgb = None
            return p

        def _add_para(text: str):
            # normal paragraph
            return doc.add_paragraph(text)

        def _add_quote(text: str):
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Pt(12)
            for r in p.runs:
                r.italic = True
            return p

        def _add_code_block(text: str):
            # Use monospace and preserve line breaks via separate runs and paragraphs
            text = _soft_wrap_long_tokens(text)
            for i, line in enumerate(text.split("\n")):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = "Courier New"
                # ensure Word uses the specified font for East Asia as well
                try:
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
                except Exception:
                    pass
                run.font.size = Pt(9)
            return True

        def _add_heading(level: int, text: str):
            lvl = min(max(level, 2), 4)  # map to Word heading levels 2-4
            h = doc.add_heading(_soft_wrap_long_tokens(text), level=lvl)
            return h

        def _add_table(rows: list[list[str]]):
            if not rows:
                return
            hdr = rows[0]
            body = rows[1:] if len(rows) > 1 else []
            table = doc.add_table(rows=len(rows), cols=len(hdr))
            table.style = "Table Grid"
            # header
            for j, val in enumerate(hdr):
                table.rows[0].cells[j].text = str(val)
            # body
            for i, row in enumerate(body, start=1):
                for j, val in enumerate(row):
                    table.rows[i].cells[j].text = str(val)
            doc.add_paragraph("")

        def _is_table_separator(line: str) -> bool:
            l = line.strip()
            if set(l) <= set("|-: ") and "-" in l and "|" in l:
                return True
            return False

        def _parse_md_table(lines: list[str], i: int):
            if i + 1 >= len(lines):
                return None, i
            header_line = lines[i]
            sep_line = lines[i + 1]
            if "|" not in header_line or not _is_table_separator(sep_line):
                return None, i

            def split_row(s: str) -> list[str]:
                cells = [c.strip() for c in s.strip().strip("|").split("|")]
                return cells

            header = split_row(header_line)
            rows = []
            j = i + 2
            while (
                j < len(lines)
                and "|" in lines[j]
                and not lines[j].strip().startswith("#")
            ):
                if lines[j].strip() == "":
                    break
                rows.append(split_row(lines[j]))
                j += 1
            width = len(header)
            header = header + [""] * (width - len(header))
            norm_rows = []
            for r in rows:
                r2 = r + [""] * (width - len(r))
                norm_rows.append(r2[:width])
            return [header] + norm_rows, j - 1

        header = _chat_export_header(chat, msgs, app_name, export_language, chat_title)
        _add_meta(header[0])
        _add_meta(header[1])
        _add_title(header[2])
        _add_meta(header[3])
        doc.add_paragraph("")

        for m in msgs:
            role = (m.get("role") or "").lower()
            content = _message_display_text_or_placeholder(m, export_language)
            _add_heading(2, _message_heading(m, export_language))

            # Demote headings by two
            demoted_lines = []
            for ln in str(content).splitlines():
                m_h = _re.match(r"^(#{1,6})\s+(.*)$", ln)
                if m_h:
                    cur = len(m_h.group(1))
                    new = min(6, cur + 2)
                    demoted_lines.append("#" * new + " " + m_h.group(2))
                else:
                    demoted_lines.append(ln)

            lines = demoted_lines
            i = 0
            in_code = False
            code_buffer = []
            para_buffer = []

            def flush_paragraph_buffer():
                if para_buffer:
                    txt = _soft_wrap_long_tokens("\n".join(para_buffer))
                    _add_para(txt)
                    doc.add_paragraph("")
                    para_buffer.clear()

            def flush_code_buffer():
                if code_buffer:
                    _add_code_block("\n".join(code_buffer))
                    doc.add_paragraph("")
                    code_buffer.clear()

            while i < len(lines):
                ln = lines[i]
                if ln.strip().startswith("```"):
                    if in_code:
                        in_code = False
                        flush_code_buffer()
                    else:
                        flush_paragraph_buffer()
                        in_code = True
                    i += 1
                    continue

                if in_code:
                    code_buffer.append(ln)
                    i += 1
                    continue

                # Table
                tbl, j = _parse_md_table(lines, i)
                if tbl:
                    flush_paragraph_buffer()
                    _add_table(tbl)
                    i = j + 1
                    continue

                # Heading
                m_h = _re.match(r"^(#{1,6})\s+(.*)$", ln)
                if m_h:
                    flush_paragraph_buffer()
                    level = len(m_h.group(1))
                    _add_heading(level, m_h.group(2))
                    i += 1
                    continue

                # Blockquote
                if ln.startswith("> "):
                    flush_paragraph_buffer()
                    _add_quote(ln[2:])
                    i += 1
                    continue

                # Unordered list
                if ln.lstrip().startswith("- ") or ln.lstrip().startswith("* "):
                    item = ln.lstrip()[2:]
                    para_buffer.append(f"• {item}")
                    i += 1
                    continue

                # Ordered list
                m_li = _re.match(r"^\s*(\d+)\.\s+(.*)$", ln)
                if m_li:
                    para_buffer.append(f"{m_li.group(1)}. {m_li.group(2)}")
                    i += 1
                    continue

                # Blank line
                if ln.strip() == "":
                    flush_paragraph_buffer()
                    i += 1
                    continue

                # Normal paragraph text
                para_buffer.append(ln)
                i += 1

            flush_code_buffer()
            flush_paragraph_buffer()

            # Attachments
            attach = _message_attachment_labels(m)
            if attach:
                for field, lst in attach.items():
                    doc.add_paragraph(
                        f"{_attachment_heading(export_language, field)}: {', '.join(lst)}"
                    )
                doc.add_paragraph("")

        if compliance_watermark:
            # Keep the compliance marker as a normal document paragraph so it
            # remains visible in Word without altering the message content.
            watermark_paragraph = doc.add_paragraph(compliance_watermark)
            for run in watermark_paragraph.runs:
                run.italic = True

        buf = BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()
        buf.close()
        return {"filename": f"{base}.docx", "type": "docx", "content": docx_bytes}

    if f == "pdf":
        # Render a polished transcript PDF with branding and basic Markdown support.
        try:
            from io import BytesIO
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import mm
            from reportlab.lib import colors
            from reportlab.platypus import Image as ReportLabImage
            from reportlab.platypus import (
                SimpleDocTemplate,
                Paragraph,
                Spacer,
                Table,
                TableStyle,
            )
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        except Exception:
            raise HTTPException(
                status_code=501,
                detail="PDF export requires the 'reportlab' package. Please install it (e.g., pip install reportlab).",
            )

        import re as _re

        chat = data.get("chat", {})
        msgs = data.get("messages", [])
        app_name = _resolve_application_name(db)
        export_language = _resolve_pdf_export_language(user_id, db)
        base_font = _register_pdf_export_font(export_language)
        chat_title = str(chat.get("title") or base or "Chat").strip() or "Chat"
        exported_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

        buf = BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=18 * mm,
            rightMargin=18 * mm,
            topMargin=18 * mm,
            bottomMargin=18 * mm,
            title=chat_title,
            author=app_name,
        )
        available_width = A4[0] - (36 * mm)

        styles = getSampleStyleSheet()
        styles.add(
            ParagraphStyle(
                name="PdfAppName",
                parent=styles["BodyText"],
                fontName=base_font,
                fontSize=13,
                leading=16,
                textColor=colors.HexColor("#111827"),
                spaceAfter=0,
            )
        )
        styles.add(
            ParagraphStyle(
                name="PdfExportMeta",
                parent=styles["BodyText"],
                fontName=base_font,
                fontSize=8.5,
                leading=11,
                textColor=colors.HexColor("#6B7280"),
            )
        )
        styles.add(
            ParagraphStyle(
                name="PdfChatTitle",
                parent=styles["Title"],
                fontName=base_font,
                fontSize=22,
                leading=26,
                spaceAfter=4,
                alignment=0,
                textColor=colors.HexColor("#111827"),
            )
        )
        styles.add(
            ParagraphStyle(
                name="PdfSubtitle",
                parent=styles["BodyText"],
                fontName=base_font,
                fontSize=9.5,
                leading=13,
                textColor=colors.HexColor("#6B7280"),
                spaceAfter=12,
            )
        )
        styles.add(
            ParagraphStyle(
                name="PdfComplianceWatermark",
                parent=styles["BodyText"],
                fontName=base_font,
                fontSize=8.5,
                leading=11,
                textColor=colors.HexColor("#6B7280"),
                spaceBefore=14,
                spaceAfter=4,
            )
        )
        styles.add(
            ParagraphStyle(
                name="PdfRoleUser",
                parent=styles["BodyText"],
                fontName=base_font,
                fontSize=10,
                leading=12,
                textColor=colors.HexColor("#0F5132"),
                spaceBefore=8,
                spaceAfter=2,
            )
        )
        styles.add(
            ParagraphStyle(
                name="PdfRoleAssistant",
                parent=styles["BodyText"],
                fontName=base_font,
                fontSize=10,
                leading=12,
                textColor=colors.HexColor("#1F2937"),
                spaceBefore=8,
                spaceAfter=2,
            )
        )
        styles.add(
            ParagraphStyle(
                name="PdfRoleTool",
                parent=styles["BodyText"],
                fontName=base_font,
                fontSize=10,
                leading=12,
                textColor=colors.HexColor("#7C2D12"),
                spaceBefore=8,
                spaceAfter=2,
            )
        )
        styles.add(
            ParagraphStyle(
                name="PdfHeading2",
                parent=styles["Heading2"],
                fontName=base_font,
                fontSize=14,
                leading=17,
                spaceBefore=8,
                spaceAfter=4,
                textColor=colors.HexColor("#111827"),
            )
        )
        styles.add(
            ParagraphStyle(
                name="PdfHeading3",
                parent=styles["Heading3"],
                fontName=base_font,
                fontSize=12.5,
                leading=15,
                spaceBefore=6,
                spaceAfter=3,
                textColor=colors.HexColor("#111827"),
            )
        )
        styles.add(
            ParagraphStyle(
                name="PdfBody",
                parent=styles["BodyText"],
                fontName=base_font,
                fontSize=10.5,
                leading=14,
                textColor=colors.HexColor("#1F2937"),
            )
        )
        styles.add(
            ParagraphStyle(
                name="PdfQuote",
                parent=styles["BodyText"],
                fontName=base_font,
                leftIndent=10,
                textColor=colors.HexColor("#4B5563"),
                borderColor=colors.HexColor("#D1D5DB"),
                borderWidth=0,
                borderPadding=5,
            )
        )
        styles.add(
            ParagraphStyle(
                name="PdfAttachment",
                parent=styles["BodyText"],
                fontName=base_font,
                textColor=colors.HexColor("#4B5563"),
                fontSize=9.3,
                leading=12,
            )
        )
        styles.add(
            ParagraphStyle(
                name="PdfAttachmentHeading",
                parent=styles["BodyText"],
                fontName=base_font,
                textColor=colors.HexColor("#374151"),
                fontSize=9.8,
                leading=12,
                spaceBefore=4,
                spaceAfter=4,
            )
        )
        styles.add(
            ParagraphStyle(
                name="PdfImageCaption",
                parent=styles["BodyText"],
                fontName=base_font,
                textColor=colors.HexColor("#6B7280"),
                fontSize=8.8,
                leading=11,
                spaceAfter=6,
            )
        )
        styles.add(
            ParagraphStyle(
                name="PdfFileBadge",
                parent=styles["BodyText"],
                fontName=base_font,
                textColor=colors.HexColor("#111827"),
                fontSize=7.5,
                leading=9,
                alignment=1,
            )
        )
        styles.add(
            ParagraphStyle(
                name="PdfFileName",
                parent=styles["BodyText"],
                fontName=base_font,
                textColor=colors.HexColor("#111827"),
                fontSize=9.5,
                leading=12,
            )
        )
        styles.add(
            ParagraphStyle(
                name="PdfFileMeta",
                parent=styles["BodyText"],
                fontName=base_font,
                textColor=colors.HexColor("#6B7280"),
                fontSize=8.3,
                leading=10,
            )
        )
        mono = ParagraphStyle(
            name="PdfCode",
            parent=styles["BodyText"],
            fontName="Courier",
            fontSize=9,
            leading=11,
            backColor=colors.HexColor("#F3F4F6"),
            textColor=colors.HexColor("#111827"),
            leftIndent=4,
            rightIndent=4,
            spaceBefore=2,
            spaceAfter=4,
            wordWrap="CJK",  # allow breaking long tokens to avoid right-side clipping
        )

        def _escape(text: str) -> str:
            # Minimal escaping for ReportLab Paragraph XML
            return (
                str(text)
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )

        def _md_inline(text: str) -> str:
            # Convert simple markdown inline to ReportLab-safe inline tags
            # Order of operations matters to avoid invalid nesting for <b>/<i>.
            s = _escape(str(text))

            # 1) Protect inline code spans first using placeholders
            code_spans: list[str] = []

            def _code_repl(m: _re.Match) -> str:
                content = m.group(1)
                token = f"@@CODE{len(code_spans)}@@"
                code_spans.append(f"<font face='Courier'>{content}</font>")
                return token

            s = _re.sub(r"`([^`]+)`", _code_repl, s)

            # 2) Bold+Italic: ***text*** or ___text___ -> <b><i>text</i></b>
            s = _re.sub(r"(?<!\\)(\*{3}|_{3})(.+?)(?<!\\)\1", r"<b><i>\2</i></b>", s)

            # 3) Bold: **text** or __text__
            s = _re.sub(r"(?<!\\)(\*{2}|_{2})(.+?)(?<!\\)\1", r"<b>\2</b>", s)

            # 4) Italic: *text* or _text_
            # Use patterns that avoid consuming ** or __ cases
            s = _re.sub(
                r"(?<!\\)(?<!\*)\*(?!\*)(.+?)(?<!\\)(?<!\*)\*(?!\*)", r"<i>\1</i>", s
            )
            s = _re.sub(r"(?<!\\)(?<!_)_(?!_)(.+?)(?<!\\)(?<!_)_(?!_)", r"<i>\1</i>", s)

            # 5) Restore code spans
            for i, tag in enumerate(code_spans):
                s = s.replace(f"@@CODE{i}@@", tag)

            return s

        def _inline_paragraph_text(text: str) -> str:
            return _md_inline(_soft_wrap_long_tokens(str(text))).replace("\n", "<br/>")

        def _soft_wrap_long_tokens(
            text: str, max_token_len: int = 60, insert_every: int = 12
        ) -> str:
            """Insert zero-width spaces into very long tokens to allow wrapping.
            Preserves whitespace and newlines; apply on a per-line basis before replacing newlines with <br/>.
            """
            ZWSP = "\u200b"
            out_lines = []
            for line in text.split("\n"):
                parts = []
                for tok in _re.split(r"(\s+)", line):
                    if not tok or tok.isspace():
                        parts.append(tok)
                        continue
                    if len(tok) > max_token_len:
                        # insert soft breaks every N chars
                        chunks = [
                            tok[i : i + insert_every]
                            for i in range(0, len(tok), insert_every)
                        ]
                        parts.append(ZWSP.join(chunks))
                    else:
                        parts.append(tok)
                out_lines.append("".join(parts))
            return "\n".join(out_lines)

        def _thinking_paragraph(text: str):
            # Escape, soft-wrap long tokens, convert newlines to <br/>, then render as Paragraph with monospace style.
            safe = _escape(text)
            safe = _soft_wrap_long_tokens(safe)
            safe = safe.replace("\n", "<br/>")
            return Paragraph(safe, mono)

        def _role_label(role: str, name: str | None = None) -> str:
            role = (role or "").lower()
            if role == "user":
                return _pdf_export_t(export_language, "pdf_export_user")
            if role == "assistant":
                return _pdf_export_t(export_language, "pdf_export_assistant")
            if role == "tool":
                return (
                    _pdf_export_t(export_language, "pdf_export_tool_named", name=name)
                    if name
                    else _pdf_export_t(export_language, "pdf_export_tool")
                )
            if role == "system":
                return _pdf_export_t(export_language, "pdf_export_system")
            return role.capitalize() or _pdf_export_t(
                export_language, "pdf_export_message"
            )

        def _role_style(role: str):
            role = (role or "").lower()
            if role == "user":
                return styles["PdfRoleUser"]
            if role == "tool":
                return styles["PdfRoleTool"]
            return styles["PdfRoleAssistant"]

        def _attachment_meta_label(attachment: PdfExportAttachment) -> str:
            parts = []
            extension = _pdf_file_extension_label(
                attachment.label, attachment.file_type
            )
            if extension:
                parts.append(extension)
            if attachment.file_size > 0:
                parts.append(_format_pdf_file_size(attachment.file_size))
            if not parts and attachment.file_type:
                parts.append(attachment.file_type)
            return " - ".join(parts) if parts else attachment.file_id

        def _attachment_file_card(attachment: PdfExportAttachment):
            extension = _pdf_file_extension_label(
                attachment.label, attachment.file_type
            )
            badge = Table(
                [[Paragraph(_escape(extension), styles["PdfFileBadge"])]],
                colWidths=[13 * mm],
                rowHeights=[13 * mm],
            )
            badge.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F3F4F6")),
                        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#D1D5DB")),
                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 1),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 1),
                        ("TOPPADDING", (0, 0), (-1, -1), 1),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                    ]
                )
            )

            card = Table(
                [
                    [
                        badge,
                        [
                            Paragraph(
                                _inline_paragraph_text(attachment.label),
                                styles["PdfFileName"],
                            ),
                            Paragraph(
                                _escape(_attachment_meta_label(attachment)),
                                styles["PdfFileMeta"],
                            ),
                        ],
                    ]
                ],
                colWidths=[16 * mm, available_width - (16 * mm)],
            )
            card.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FFFFFF")),
                        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#E5E7EB")),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            return card

        def _attachment_image_flowable(attachment: PdfExportAttachment):
            if not attachment.path or not attachment.file_type.startswith("image/"):
                return None
            try:
                from PIL import Image as PilImage

                image_source = str(attachment.path)
                pil_source = attachment.path
                if (
                    attachment.file_type == "image/svg+xml"
                    or attachment.path.suffix.lower() == ".svg"
                ):
                    png_bytes = rasterize_svg_to_png_bytes(
                        svg_path=attachment.path,
                        output_width=1600,
                        output_height=1600,
                    )
                    image_source = BytesIO(png_bytes)
                    pil_source = BytesIO(png_bytes)

                with PilImage.open(pil_source) as image:
                    width, height = image.size
                if width <= 0 or height <= 0:
                    return None

                max_image_height = 130 * mm
                scale = min(
                    available_width / float(width),
                    max_image_height / float(height),
                    1.0,
                )
                display_width = float(width) * scale
                display_height = float(height) * scale
                return ReportLabImage(
                    image_source,
                    width=display_width,
                    height=display_height,
                    kind="proportional",
                )
            except Exception:
                logger.debug(
                    "Unable to render chat PDF attachment image", exc_info=True
                )
                return None

        def _append_pdf_attachments(story: list, message: dict) -> None:
            attachments = _resolve_pdf_export_attachments(message, db, str(user_id))
            if not attachments:
                return

            for field in ATTACHMENT_FIELDS:
                field_attachments = attachments.get(field) or []
                if not field_attachments:
                    continue
                story.append(
                    Paragraph(
                        _escape(_attachment_heading(export_language, field)),
                        styles["PdfAttachmentHeading"],
                    )
                )
                for attachment in field_attachments:
                    image_flowable = _attachment_image_flowable(attachment)
                    if image_flowable is not None:
                        story.append(image_flowable)
                        story.append(
                            Paragraph(
                                _escape(
                                    f"{attachment.label} - {_attachment_meta_label(attachment)}"
                                ),
                                styles["PdfImageCaption"],
                            )
                        )
                    else:
                        story.append(_attachment_file_card(attachment))
                        story.append(Spacer(1, 5))
                story.append(Spacer(1, 2))

        def _add_brand_header(story: list) -> None:
            icon_path = _resolve_pdf_icon_path()
            icon_flowable = None
            if icon_path:
                try:
                    if icon_path.suffix.lower() == ".svg":
                        png_bytes = rasterize_svg_to_png_bytes(
                            svg_path=icon_path,
                            output_width=96,
                            output_height=96,
                        )
                        icon_flowable = ReportLabImage(
                            BytesIO(png_bytes),
                            width=11 * mm,
                            height=11 * mm,
                            kind="proportional",
                        )
                    else:
                        icon_flowable = ReportLabImage(
                            str(icon_path),
                            width=11 * mm,
                            height=11 * mm,
                            kind="proportional",
                        )
                except Exception:
                    icon_flowable = None

            if icon_flowable is None:
                initial = _escape(app_name[:1].upper() or "C")
                icon_flowable = Table(
                    [[Paragraph(initial, styles["PdfAppName"])]],
                    colWidths=[11 * mm],
                    rowHeights=[11 * mm],
                )
                icon_flowable.setStyle(
                    TableStyle(
                        [
                            (
                                "BACKGROUND",
                                (0, 0),
                                (-1, -1),
                                colors.HexColor("#EEF2FF"),
                            ),
                            ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#C7D2FE")),
                            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ]
                    )
                )

            header = Table(
                [
                    [
                        icon_flowable,
                        [
                            Paragraph(_escape(app_name), styles["PdfAppName"]),
                            Paragraph(
                                f"{_escape(_pdf_export_t(export_language, 'pdf_export_title'))} - "
                                f"{_escape(_pdf_export_t(export_language, 'pdf_export_generated', date=exported_at))}",
                                styles["PdfExportMeta"],
                            ),
                        ],
                    ]
                ],
                colWidths=[14 * mm, available_width - (14 * mm)],
            )
            header.setStyle(
                TableStyle(
                    [
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                        (
                            "LINEBELOW",
                            (0, 0),
                            (-1, -1),
                            0.6,
                            colors.HexColor("#E5E7EB"),
                        ),
                    ]
                )
            )
            story.append(header)
            story.append(Spacer(1, 10))
            story.append(
                Paragraph(_inline_paragraph_text(chat_title), styles["PdfChatTitle"])
            )

            created = _format_export_datetime(chat.get("created_at"))
            updated = _format_export_datetime(chat.get("last_updated_at"))
            count_key = (
                "pdf_export_message_count_one"
                if len(msgs) == 1
                else "pdf_export_message_count_other"
            )
            details = [_pdf_export_t(export_language, count_key, count=len(msgs))]
            if created:
                details.append(
                    _pdf_export_t(export_language, "pdf_export_started", date=created)
                )
            if updated:
                details.append(
                    _pdf_export_t(export_language, "pdf_export_updated", date=updated)
                )
            story.append(Paragraph(_escape(" | ".join(details)), styles["PdfSubtitle"]))

        def _page_footer(canvas, document):
            canvas.saveState()
            canvas.setStrokeColor(colors.HexColor("#E5E7EB"))
            canvas.setLineWidth(0.5)
            canvas.line(
                document.leftMargin, 12 * mm, A4[0] - document.rightMargin, 12 * mm
            )
            canvas.setFillColor(colors.HexColor("#6B7280"))
            canvas.setFont(base_font, 8)
            canvas.drawString(document.leftMargin, 8 * mm, app_name)
            canvas.drawRightString(
                A4[0] - document.rightMargin,
                8 * mm,
                _pdf_export_t(export_language, "pdf_export_page", page=document.page),
            )
            canvas.restoreState()

        def _is_table_separator(line: str) -> bool:
            l = line.strip()
            if set(l) <= set("|-: ") and "-" in l and "|" in l:
                return True
            return False

        def _parse_md_table(lines: list[str], i: int):
            # Expect header, separator, then rows starting with '|'
            if i + 1 >= len(lines):
                return None, i
            header_line = lines[i]
            sep_line = lines[i + 1]
            if "|" not in header_line or not _is_table_separator(sep_line):
                return None, i

            def split_row(s: str) -> list[str]:
                cells = [c.strip() for c in s.strip().strip("|").split("|")]
                return cells

            header = split_row(header_line)
            rows = []
            j = i + 2
            while (
                j < len(lines)
                and "|" in lines[j]
                and not lines[j].strip().startswith("#")
            ):
                if lines[j].strip() == "":
                    break
                rows.append(split_row(lines[j]))
                j += 1
            # Normalize row lengths
            width = len(header)
            header = header + [""] * (width - len(header))
            norm_rows = []
            for r in rows:
                r2 = r + [""] * (width - len(r))
                norm_rows.append(r2[:width])
            return [header] + norm_rows, j - 1

        story = []
        _add_brand_header(story)

        for m in msgs:
            role = (m.get("role") or "").lower()
            content = _message_display_content(m)
            name = m.get("name")
            created_at = _format_export_datetime(m.get("created_at"))

            label = _role_label(role, name)
            if created_at:
                label = f"{label} - {created_at}"
            story.append(Paragraph(_escape(label), _role_style(role)))

            if role == "tool" and not content:
                content = _pdf_export_t(export_language, "pdf_export_tool_activity")

            # Demote headings by two to keep document hierarchy
            demoted_lines = []
            for ln in str(content).splitlines():
                m_h = _re.match(r"^(#{1,6})\s+(.*)$", ln)
                if m_h:
                    cur = len(m_h.group(1))
                    new = min(6, cur + 2)
                    demoted_lines.append("#" * new + " " + m_h.group(2))
                else:
                    demoted_lines.append(ln)

            lines = demoted_lines
            i = 0
            in_code = False
            code_buffer = []
            para_buffer = []

            def flush_paragraph_buffer():
                if para_buffer:
                    paragraph_text = _inline_paragraph_text("\n".join(para_buffer))
                    story.append(Paragraph(paragraph_text, styles["PdfBody"]))
                    story.append(Spacer(1, 4))
                    para_buffer.clear()

            def flush_code_buffer():
                if code_buffer:
                    story.append(_thinking_paragraph("\n".join(code_buffer)))
                    story.append(Spacer(1, 4))
                    code_buffer.clear()

            while i < len(lines):
                ln = lines[i]
                if ln.strip().startswith("```"):
                    if in_code:
                        # close code
                        in_code = False
                        flush_code_buffer()
                    else:
                        # open code
                        flush_paragraph_buffer()
                        in_code = True
                    i += 1
                    continue

                if in_code:
                    code_buffer.append(ln)
                    i += 1
                    continue

                # Try table
                tbl, j = _parse_md_table(lines, i)
                if tbl:
                    flush_paragraph_buffer()
                    # Build table
                    col_count = max(len(tbl[0]), 1)
                    col_width = available_width / col_count
                    data_rows = [
                        [
                            Paragraph(_inline_paragraph_text(col), styles["PdfBody"])
                            for col in row
                        ]
                        for row in tbl
                    ]
                    t = Table(
                        data_rows, colWidths=[col_width] * col_count, hAlign="LEFT"
                    )
                    t.setStyle(
                        TableStyle(
                            [
                                ("FONTNAME", (0, 0), (-1, 0), base_font),
                                (
                                    "BACKGROUND",
                                    (0, 0),
                                    (-1, 0),
                                    colors.HexColor("#F3F4F6"),
                                ),
                                (
                                    "TEXTCOLOR",
                                    (0, 0),
                                    (-1, 0),
                                    colors.HexColor("#222222"),
                                ),
                                (
                                    "GRID",
                                    (0, 0),
                                    (-1, -1),
                                    0.5,
                                    colors.HexColor("#D1D5DB"),
                                ),
                                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                                ("TOPPADDING", (0, 0), (-1, -1), 4),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                            ]
                        )
                    )
                    story.append(t)
                    story.append(Spacer(1, 8))
                    i = j + 1
                    continue

                # Headings
                m_h = _re.match(r"^(#{1,6})\s+(.*)$", ln)
                if m_h:
                    flush_paragraph_buffer()
                    level = len(m_h.group(1))
                    text = _inline_paragraph_text(m_h.group(2))
                    if level <= 2:
                        story.append(Paragraph(text, styles["PdfHeading2"]))
                    elif level == 3:
                        story.append(Paragraph(text, styles["PdfHeading3"]))
                    else:
                        story.append(Paragraph(text, styles["PdfBody"]))
                    i += 1
                    continue

                # Blockquote
                if ln.startswith("> "):
                    flush_paragraph_buffer()
                    story.append(
                        Paragraph(_inline_paragraph_text(ln[2:]), styles["PdfQuote"])
                    )
                    story.append(Spacer(1, 4))
                    i += 1
                    continue

                # Unordered list
                if ln.lstrip().startswith("- ") or ln.lstrip().startswith("* "):
                    item = ln.lstrip()[2:]
                    para_buffer.append(f"• {item}")
                    i += 1
                    continue

                # Ordered list
                m_li = _re.match(r"^\s*(\d+)\.\s+(.*)$", ln)
                if m_li:
                    para_buffer.append(f"{m_li.group(1)}. {m_li.group(2)}")
                    i += 1
                    continue

                # Blank line -> flush paragraph
                if ln.strip() == "":
                    flush_paragraph_buffer()
                    i += 1
                    continue

                # Normal paragraph text
                para_buffer.append(ln)
                i += 1

            # Flush remaining buffers for this message
            flush_code_buffer()
            flush_paragraph_buffer()

            _append_pdf_attachments(story, m)

        if compliance_watermark:
            # PDF exports need a real flowable rather than a raw string so the
            # marker is visible and wrapped correctly on narrow pages.
            story.append(
                Paragraph(
                    _inline_paragraph_text(compliance_watermark),
                    styles["PdfComplianceWatermark"],
                )
            )

        # Build PDF
        doc.build(story, onFirstPage=_page_footer, onLaterPages=_page_footer)
        pdf_bytes = buf.getvalue()
        buf.close()
        return {"filename": f"{base}.pdf", "type": "pdf", "content": pdf_bytes}

    raise HTTPException(
        status_code=400, detail="Unsupported format. Use 'json', 'txt', 'pdf', or 'md'."
    )


# -------------------
# Export Chat Full
# -------------------
def _serialize_chat_message_export(message: ChatMessages) -> dict[str, Any]:
    attachments = _collect_message_attachments(message)
    decoded_content = _decode_json_content(getattr(message, "content", None))
    blocks = (
        decoded_content
        if isinstance(decoded_content, list)
        else [decoded_content]
        if isinstance(decoded_content, dict)
        else []
    )
    first_block = (
        next((block for block in blocks if isinstance(block, dict)), {})
        if blocks
        else {}
    )
    return {
        "id": message.id,
        "model_id": message.model_id,
        "role": message.role,
        "content": message.content,
        "reference_id": getattr(message, "reference_id", None),
        "generation": deepcopy(message.generation)
        if isinstance(message.generation, dict)
        else message.generation,
        "thinking": getattr(message, "thinking", None),
        "retry_count": getattr(message, "retry_count", 0),
        "bookmarked": bool(getattr(message, "bookmarked", False)),
        "images": attachments["images"],
        "videos": attachments["videos"],
        "audios": attachments["audios"],
        "documents": attachments["documents"],
        "youtube": first_block.get("youtube"),
        "sources": first_block.get("sources"),
        "name": first_block.get("tool_name"),
        "meta": first_block.get("meta"),
        "created_at": message.created_at.isoformat()
        if getattr(message, "created_at", None)
        else None,
    }


def _build_chat_export_payload(
    chat: Chats, db, *, include_attention: bool = True
) -> dict[str, Any]:
    messages = (
        db.query(ChatMessages)
        .filter(ChatMessages.chat_id == chat.id)
        .order_by(ChatMessages.created_at.asc(), ChatMessages.id.asc())
    )

    read_version = 0
    if include_attention:
        receipt = (
            db.query(ChatReadState)
            .filter(
                ChatReadState.user_id == chat.user_id, ChatReadState.chat_id == chat.id
            )
            .first()
        )
        read_version = int(receipt.read_response_version or 0) if receipt else 0
    chat_obj = {
        "id": chat.id,
        "user_id": chat.user_id,
        "title": chat.title,
        "project_id": chat.project_id,
        "share": sanitize_chat_share_for_export(chat.share),
        "archived": chat.archived,
        "pinned_position": chat.pinned_position,
        "meta": deepcopy(chat.meta) if isinstance(chat.meta, dict) else chat.meta,
        "created_at": chat.created_at.isoformat()
        if getattr(chat, "created_at", None)
        else None,
        "last_updated_at": chat.last_updated_at.isoformat()
        if getattr(chat, "last_updated_at", None)
        else None,
        # Export the portable behavior rather than leaking internal counters.
    }
    if include_attention:
        chat_obj["has_unread_response"] = (
            int(getattr(chat, "response_version", 0) or 0) > read_version
        )

    return {
        "chat": chat_obj,
        "messages": [
            _serialize_chat_message_export(message)
            for message in _iter_export_query_rows(messages)
        ],
        "deep_research_runs": _export_deep_research_runs_for_chat(
            chat.user_id, chat.id, db
        ),
    }


def export_chat_full(
    user_id: str,
    chat_id: str,
    db,
    include_deleted_or_temp: bool = False,
    compliance_watermark: str = "",
):
    """
    Build a complete export payload for a chat owned by user_id.

    Returns a dict with keys:
      - chat: {id, user_id, title, project_id, share, archived, pinned, created_at, last_updated_at}
      - messages: [ {id, role, content, thinking, images, videos, audios, documents, name, meta, created_at} ]
    Messages are ordered by created_at (asc), then id for stability.
    When supplied, ``compliance_watermark`` is applied to a copied message
    payload so callers can reuse the same builder for protected exports.
    """
    # Ensure the chat exists and belongs to the user
    chat = db.query(Chats).filter(Chats.id == chat_id, Chats.user_id == user_id).first()
    if not chat:
        raise HTTPException(status_code=404, detail="Chat not found!")
    if not include_deleted_or_temp and is_chat_excluded_from_default_export(chat):
        raise HTTPException(status_code=404, detail="Chat not found!")
    payload = _build_chat_export_payload(chat, db)
    return apply_compliance_watermark_to_chat_export(payload, compliance_watermark)


def get_all_chats_export_audit_details(
    db: Session,
    include_deleted_or_temp: bool = False,
) -> dict[str, Any]:
    """Return audit details before streaming an all-users chat export."""
    user_reference_map: dict[str, str] = {}
    exported_count = 0

    chats_query = (
        db.query(Chats, User.email)
        .outerjoin(User, User.id == Chats.user_id)
        .order_by(Chats.created_at.asc(), Chats.id.asc())
    )
    for row in _iter_export_query_rows(chats_query):
        try:
            chat, user_email = row
        except (TypeError, ValueError):
            chat, user_email = row, None
        if not include_deleted_or_temp and is_chat_excluded_from_default_export(chat):
            continue

        reference_token = build_email_reference_token(user_email)
        if reference_token:
            user_reference_map[str(chat.user_id)] = reference_token
        exported_count += 1

    return {
        "count": exported_count,
        "user_reference_count": len(user_reference_map),
        "include_deleted_or_temp": bool(include_deleted_or_temp),
    }


def iter_all_chats_export_json(
    db: Session,
    include_deleted_or_temp: bool = False,
    *,
    counters: dict[str, Any] | None = None,
) -> Iterator[str]:
    """Stream an all-users chat export JSON document without materializing the payload."""
    generated_at = datetime.now(timezone.utc).isoformat()
    user_reference_map: dict[str, str] = {}
    exported_count = 0
    watermark_resolver = ComplianceWatermarkResolver(db)

    yield "{"
    yield f'"export_type":{_json_dumps("chats")},'
    yield f'"export_version":{_json_dumps(current_chat_export_version)},'
    yield f'"generated_at":{_json_dumps(generated_at)},'
    yield '"data":{"chats":['

    first_chat = True
    chats_query = (
        db.query(Chats, User.email)
        .outerjoin(User, User.id == Chats.user_id)
        .order_by(Chats.created_at.asc(), Chats.id.asc())
    )
    for row in _iter_export_query_rows(chats_query):
        try:
            chat, user_email = row
        except (TypeError, ValueError):
            chat, user_email = row, None
        if not include_deleted_or_temp and is_chat_excluded_from_default_export(chat):
            continue

        reference_token = build_email_reference_token(user_email)
        if reference_token:
            user_reference_map[str(chat.user_id)] = reference_token

        if not first_chat:
            yield ","
        first_chat = False
        # The streamed admin export remains content-focused and avoids an
        # additional per-chat receipt query. Full database backups and user
        # exports retain attention state.
        chat_payload = _build_chat_export_payload(chat, db, include_attention=False)
        apply_compliance_watermark_to_chat_export(
            chat_payload,
            watermark_resolver.for_user(chat.user_id),
        )
        yield _json_dumps(chat_payload)
        exported_count += 1

    yield "],"
    yield f'"count":{exported_count},'
    yield f'"user_reference_map":{_json_dumps(user_reference_map)}'
    yield "}}"

    if counters is not None:
        counters["count"] = exported_count
        counters["user_reference_count"] = len(user_reference_map)


def iter_user_chats_export_json(
    user_id: str,
    db: Session,
    include_deleted_or_temp: bool = False,
    *,
    counters: dict[str, Any] | None = None,
) -> Iterator[str]:
    """Stream one user's versioned chat export without building a giant list.

    The selected-user export deliberately preserves attention state, matching
    the existing user-data export contract. Only one chat payload is held in
    memory at a time; messages and related run data are released after their
    JSON chunk has been written.
    """
    generated_at = datetime.now(timezone.utc).isoformat()
    exported_count = 0
    watermark_resolver = ComplianceWatermarkResolver(db)

    yield "{"
    yield f'"export_type":{_json_dumps("chats")},'
    yield f'"export_version":{_json_dumps(current_chat_export_version)},'
    yield f'"generated_at":{_json_dumps(generated_at)},'
    yield '"data":{"chats":['

    first_chat = True
    chats_query = (
        db.query(Chats)
        .filter(Chats.user_id == user_id)
        .order_by(Chats.created_at.asc(), Chats.id.asc())
    )
    for chat in _iter_export_query_rows(chats_query):
        if not include_deleted_or_temp and is_chat_excluded_from_default_export(chat):
            continue

        if not first_chat:
            yield ","
        first_chat = False
        chat_payload = _build_chat_export_payload(chat, db, include_attention=True)
        apply_compliance_watermark_to_chat_export(
            chat_payload,
            watermark_resolver.for_user(user_id),
        )
        yield _json_dumps(chat_payload)
        exported_count += 1

    yield "],"
    yield f'"count":{exported_count}'
    yield "}}"

    if counters is not None:
        counters["count"] = exported_count


def _export_deep_research_runs_for_chat(
    user_id: str,
    chat_id: str,
    db,
) -> list[dict[str, Any]]:
    """Export durable v2 run metadata and bounded canonical report files."""

    from app.tools.deep_research.models import (
        DeepResearchArtifact,
        DeepResearchRun,
    )
    from app.tools.deep_research.storage import (
        get_deep_research_run_storage_provider,
        materialize_deep_research_artifact,
    )

    runs = (
        db.query(DeepResearchRun)
        .filter(
            DeepResearchRun.user_id == str(user_id),
            DeepResearchRun.chat_id == str(chat_id),
        )
        .order_by(DeepResearchRun.created_at.asc(), DeepResearchRun.id.asc())
        .all()
    )
    exported: list[dict[str, Any]] = []
    # Keep one binary budget for the whole chat. Import validation applies its
    # size ceiling to the complete Deep Research payload, so resetting this for
    # every run could create an export that its matching importer rejects.
    exported_encoded_bytes = 0
    for run in runs:
        evidence = [dict(item) for item in run.evidence or [] if isinstance(item, dict)]
        artifacts = [
            DeepResearchArtifact.from_dict(item)
            for item in run.artifacts or []
            if isinstance(item, dict)
        ]
        file_contents: dict[str, str] = {}
        for relative_path in {
            run.final_report_path,
            run.final_html_path,
            run.manifest_path,
            "session.json",
            "citations.json",
        }:
            if not relative_path:
                continue
            try:
                path = materialize_deep_research_artifact(
                    run.user_id,
                    run.id,
                    relative_path,
                    storage_provider=get_deep_research_run_storage_provider(run),
                )
                if path.stat().st_size <= 2 * 1024 * 1024:
                    file_contents[relative_path] = path.read_text(
                        encoding="utf-8",
                        errors="replace",
                    )
            except Exception:
                logger.warning(
                    "Unable to include a Deep Research report in chat export",
                    extra={"run_id": run.id, "relative_path": relative_path},
                )

        # Chat exports remain bounded, but include the validated binary report
        # assets so imported Markdown reports do not lose charts or web images.
        artifact_contents: dict[str, dict[str, str]] = {}
        for artifact in artifacts:
            relative_path = str(artifact.relative_path or "").strip()
            if (
                not relative_path.startswith("artifacts/")
                or artifact.validation_status != "validated"
            ):
                continue
            try:
                path = materialize_deep_research_artifact(
                    run.user_id,
                    run.id,
                    relative_path,
                    storage_provider=get_deep_research_run_storage_provider(run),
                )
                size = path.stat().st_size
                if size <= 0 or size > 10 * 1024 * 1024:
                    continue
                encoded_data = base64.b64encode(path.read_bytes()).decode("ascii")
                # The matching import limit applies to serialized JSON, so the
                # budget must count base64 bytes rather than the smaller source
                # file. A final exact payload pass below also accounts for
                # metadata and JSON structure.
                encoded_size = len(encoded_data.encode("ascii"))
                if (
                    exported_encoded_bytes + encoded_size
                    > CHAT_IMPORT_MAX_DEEP_RESEARCH_BYTES_PER_CHAT
                ):
                    continue
                artifact_contents[relative_path] = {
                    "encoding": "base64",
                    "media_type": artifact.media_type
                    or mimetypes.guess_type(path.name)[0]
                    or "application/octet-stream",
                    "data": encoded_data,
                }
                exported_encoded_bytes += encoded_size
            except Exception:
                logger.warning(
                    "Unable to include a Deep Research artifact in chat export",
                    extra={"run_id": run.id, "relative_path": relative_path},
                )

        config_snapshot = deepcopy(run.config_snapshot or {})
        config_snapshot.pop("quality_profile", None)
        config_snapshot.pop("budgets", None)
        exported.append(
            {
                "id": run.id,
                "generation_id": run.generation_id,
                "query": run.query,
                "execution_mode": run.execution_mode,
                "output_format": run.output_format,
                "status": run.status,
                "phase": run.phase,
                "provider_id": run.provider_id,
                "model_id": run.model_id,
                "model_name": run.model_name,
                "prompt_version": run.prompt_version,
                "revision_round": run.revision_round,
                "max_revision_rounds": run.max_revision_rounds,
                "cancel_requested": run.cancel_requested,
                "started_at": run.started_at.isoformat() if run.started_at else None,
                "completed_at": run.completed_at.isoformat()
                if run.completed_at
                else None,
                "created_at": run.created_at.isoformat() if run.created_at else None,
                "updated_at": run.updated_at.isoformat() if run.updated_at else None,
                "final_report_path": run.final_report_path,
                "final_html_path": run.final_html_path,
                "manifest_path": run.manifest_path,
                "error_code": run.error_code,
                "error_message_key": run.error_message_key,
                "config_snapshot": config_snapshot,
                "usage": deepcopy(run.usage or {}),
                "quality_gate": deepcopy(run.quality_gate or {}),
                "result_meta": deepcopy(run.result_meta or {}),
                "file_contents": file_contents,
                "artifact_contents": artifact_contents,
                "evidence": deepcopy(evidence),
                "artifacts": [
                    {
                        "stable_id": artifact.stable_id,
                        "source_phase": artifact.source_phase,
                        "original_filename": artifact.original_filename,
                        "relative_path": artifact.relative_path,
                        "media_type": artifact.media_type,
                        "kind": artifact.kind,
                        "size_bytes": artifact.size_bytes,
                        "sha256": artifact.sha256,
                        "caption": artifact.caption,
                        "alt_text": artifact.alt_text,
                        "source_url": artifact.source_url,
                        "attribution": artifact.attribution,
                        "license_name": artifact.license_name,
                        "validation_status": artifact.validation_status,
                        "meta": deepcopy(artifact.meta or {}),
                    }
                    for artifact in artifacts
                ],
            }
        )
    return _trim_deep_research_export_to_import_limit(exported)


def _deep_research_export_size_bytes(runs: list[dict[str, Any]]) -> int:
    """Return the exact UTF-8 JSON size checked by the import schema."""

    return len(
        # Match ImportedChatEntry's validator exactly, including JSON's
        # default ASCII escaping for non-ASCII report metadata.
        json.dumps(runs, separators=(",", ":")).encode("utf-8")
    )


def _trim_deep_research_export_to_import_limit(
    runs: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Drop optional embedded files until this export is accepted on import.

    Durable run metadata remains intact. Binary artifacts are removed first,
    followed by text copies when an unusually large number of reports shares a
    chat. The importer marks omitted files as unavailable instead of rejecting
    the complete chat backup.
    """

    def remove_last_file(collection_key: str) -> bool:
        for run in reversed(runs):
            collection = run.get(collection_key)
            if not isinstance(collection, dict) or not collection:
                continue
            last_key = next(reversed(collection))
            collection.pop(last_key, None)
            return True
        return False

    while (
        _deep_research_export_size_bytes(runs)
        > CHAT_IMPORT_MAX_DEEP_RESEARCH_BYTES_PER_CHAT
    ):
        if remove_last_file("artifact_contents"):
            continue
        if remove_last_file("file_contents"):
            continue
        # Runtime metadata is bounded independently and should never reach this
        # branch. Failing explicitly is safer than emitting a backup that the
        # matching importer is guaranteed to reject.
        raise ValueError(
            "Deep Research metadata exceeds the portable chat export limit."
        )
    return runs


# -------------------
# Export All Chats
# -------------------
def export_all_chats(db: Session, include_deleted_or_temp: bool = False) -> dict:
    """Return a versioned export payload containing all chats and messages."""

    all_chats = db.query(Chats).order_by(Chats.created_at.asc()).all()
    exportable_chats = [
        chat
        for chat in all_chats
        if include_deleted_or_temp or not is_chat_excluded_from_default_export(chat)
    ]
    items: list[dict] = []
    user_reference_map: dict[str, str] = {}
    watermark_resolver = ComplianceWatermarkResolver(db)

    user_ids_in_chats = {str(chat.user_id) for chat in exportable_chats}
    for user in db.query(User).filter(User.id.in_(user_ids_in_chats)).all():
        if not getattr(user, "id", None):
            continue
        reference_token = build_email_reference_token(getattr(user, "email", None))
        if reference_token:
            user_reference_map[str(user.id)] = reference_token

    for chat in exportable_chats:
        chat_payload = export_chat_full(
            chat.user_id,
            chat.id,
            db,
            include_deleted_or_temp=include_deleted_or_temp,
            compliance_watermark=watermark_resolver.for_user(chat.user_id),
        )
        items.append(chat_payload)

    return {
        "export_type": "chats",
        "export_version": current_chat_export_version,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "data": {
            "user_reference_map": user_reference_map,
            "chats": items,
            "count": len(items),
        },
    }
